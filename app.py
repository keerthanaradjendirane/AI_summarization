import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from flask import Flask, request, send_file
import tempfile
import speech_recognition as sr
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.text_rank import TextRankSummarizer
import assemblyai as aai
import docx

app = Flask(__name__)

# Email credentials
sender_email = 'keerthanaradjendirane@gmail.com'
receiver_email = 'keerthu190805@gmail.com'
password = 'jsromkntsftugyxp'  # Update with your actual email password

# Set up AssemblyAI API key
aai.settings.api_key = "dc62769570b1402d97f6ce79c836cd4d"
transcriber = aai.Transcriber()

# Initialize the recognizer
recognizer = sr.Recognizer()

def text_summarizer(text, sentences_count=3):
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = TextRankSummarizer()
    summary = summarizer(parser.document, sentences_count)
    summarized_text = " ".join([str(sentence) for sentence in summary])
    return summarized_text

# Global variables
transcriptHistory = ''

@app.route('/')
def index():
    return '''
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Speech Summarization</title>
        <style>
        body {
            background-color: #333; /* Dark grey background color */
            color: #fff; /* White text color */
            font-family: Arial, sans-serif;
            margin: 0;
            padding: 0;
            display: flex;
            flex-direction: column;
            align-items: center;
            justify-content: center;
            height: 100vh;
        }

        h1 {
            font-size: 36px;
            margin-bottom: 20px;
            margin-top: 40px; /* Added margin-top */
        }

        .button-container {
            display: flex;
            justify-content: space-between;
            width: 100%;
            max-width: 600px;
            margin-bottom: 60px; /* Increased margin for a larger gap */
        }

        button {
            font-size: 20px;
            padding: 10px 20px;
            background-color: #8a2be2; /* Purple/Violet button background color */
            color: #fff; /* White button text color */
            border: none;
            border-radius: 5px;
            cursor: pointer;
            transition: background-color 0.3s ease;
        }

        button:hover {
            background-color: #6a1b9a; /* Darker purple/violet on hover */
        }

        /* Add margin-right to the first button for gap */
        button:first-child {
            margin-right: 50px; /* Increased margin for a larger gap */
        }

        #status {
            font-size: 24px;
            margin-bottom: 10px;
        }

        #summarizedText {
            font-size: 22px;
            margin-top: 20px;
            background-color: #333; /* Dark grey box background color */
            padding: 10px;
            border-radius: 5px;
            text-align: center;
            width: 80%; /* Adjust width as needed */
            max-width: 600px; /* Maximum width for the box */
        }

        #downloadLink {
            font-size: 20px;
            color: #8a2be2; /* Purple/Violet download link color */
            margin-top: 10px;
            text-decoration: none;
        }

        .quote-container {
            background-color: #333; /* Dark grey background color */
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 40px; /* Added margin for separation */
            width: 80%; /* Adjust width as needed */
            max-width: 600px; /* Maximum width for the container */
        }

        .quote {
            font-size: 24px;
            margin: 0;
            font-style: italic;
            text-align: center;
        }
        </style>
    </head>
    <body>
        <div class="quote-container">
            <p class="quote">
                "Simplify your thoughts. Summarize your words."
            </p>
        </div>

        <h1>Speech Summarization</h1>
        <div class="button-container">
            <button onclick="startSummarization()">Start Summarizing</button>
            <button onclick="stopSummarization()">Stop Summarizing</button>
        </div>
        <div id="status"></div>
        <div id="summarizedText"></div>
        <a id="downloadLink" class="download-link" href="/download">Download Word Document</a>

        <script>
            let recognition;
            let isSummarizing = false;
            let transcriptHistory = '';

            function startSummarization() {
                recognition = new webkitSpeechRecognition();
                recognition.continuous = true;
                recognition.interimResults = false;

                recognition.onstart = function() {
                    document.getElementById('status').innerText = 'Listening... Speak now.';
                }

                recognition.onresult = function(event) {
                    let transcript = event.results[event.results.length - 1][0].transcript;
                    transcriptHistory += transcript + ' ';
                }

                recognition.start();
                isSummarizing = true;
            }

            function stopSummarization() {
                if (recognition && isSummarizing) {
                    recognition.stop();
                    document.getElementById('status').innerText = 'Speech recognition stopped.';
                    isSummarizing = false;
                    summarizeTranscript();
                }
            }

            function summarizeTranscript() {
                if (transcriptHistory.trim() !== '') {
                    fetch('/summarize', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json'
                        },
                        body: JSON.stringify({ speech: transcriptHistory })
                    })
                    .then(response => response.json())
                    .then(data => {
                        // Not displaying on the screen, just triggering the download
                        window.location.href = '/download?summarized_text=' + data.summarized_text;
                        sendEmail(data.summarized_text); // Call the sendEmail function after download
                    })
                    .catch(error => console.error('Error:', error));
                } else {
                    document.getElementById('summarizedText').innerText = 'No speech detected.';
                }
            }

            function sendEmail(summarizedText) {
                fetch('/send-email', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json'
                    },
                    body: JSON.stringify({ summarized_text: summarizedText })
                })
                .then(response => console.log('Email sent:', response))
                .catch(error => console.error('Error sending email:', error));
            }
        </script>
    </body>
    </html>
    '''

@app.route('/summarize', methods=['POST'])
def summarize():
    global transcriptHistory
    data = request.get_json()
    speech = data['speech']

    # Perform summarization
    summarized_text = text_summarizer(speech)  # Using the new summarization logic
    
    # Reset transcriptHistory for the next session
    transcriptHistory = ''

    return {'summarized_text': summarized_text}

@app.route('/download', methods=['GET'])
def download():
    # Get summarized text from query params
    summarized_text = request.args.get('summarized_text', '')

    # Create a Word document
    doc = docx.Document()
    doc.add_paragraph(summarized_text)

    # Save the Word document to a temporary file
    temp_filename = 'summarized_text.docx'
    doc.save(temp_filename)

    return send_file(temp_filename, as_attachment=True)

@app.route('/send-email', methods=['POST'])
def send_email():
    data = request.get_json()
    summarized_text = data['summarized_text']

    # Email body
    body = 'Hello, Attached is the summarized document from the speech.'

    # Create message container
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = 'Summarized Speech Report'

    # Attach the body to the message
    msg.attach(MIMEText(body, 'plain'))

    # Create a Word document in memory
    doc = docx.Document()
    doc.add_paragraph(summarized_text)

    # Save the Word document to a temporary file in memory
    temp_file = tempfile.NamedTemporaryFile(delete=False)
    temp_filename = temp_file.name
    doc.save(temp_filename)

    # Attach the Word document to the email from the temporary file in memory
    with open(temp_filename, 'rb') as attachment:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', "attachment; filename= %s" % 'summarized_text.docx')
        msg.attach(part)

    try:
        # Connect to Gmail's SMTP server
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, password)

        # Send the email
        server.sendmail(sender_email, receiver_email, msg.as_string())

        # Quit the server
        server.quit()

        return {'message': 'Email sent successfully!'}
    except Exception as e:
        return {'error': str(e)}

if __name__ == '__main__':
    app.run(debug=True)
